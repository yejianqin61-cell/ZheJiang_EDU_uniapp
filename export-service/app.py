"""
AI智能组卷 — 文档导出服务
POST /export/docx  → 生成 DOCX
POST /export/pdf   → 生成 PDF (需 LibreOffice)
"""

import io
import os
import subprocess
import tempfile
from flask import Flask, request, jsonify, send_file

app = Flask(__name__)

# ── DOCX Generation ──────────────────────────────────────────

@app.route("/export/docx", methods=["POST"])
def export_docx():
    """
    Request JSON:
    {
      "title": "试卷标题",
      "questions": [
        {
          "index": 1,
          "type": "single_choice",
          "content": "题目正文",
          "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
          "answer": "A",
          "analysis": "解析内容",
          "score": 5
        }
      ]
    }
    """
    try:
        data = request.get_json()
        title = data.get("title", "试卷")
        questions = data.get("questions", [])

        doc_bytes = generate_docx(title, questions)

        filename = f"{title}.docx"
        return send_file(
            io.BytesIO(doc_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _render_content_with_images(doc, content: str, images: list):
    """
    渲染题目内容，将 Markdown ![](url) 替换为插入的图片。

    逻辑：
    1. 按 ![](url) 分割 content，文本段→段落，图片段→插入图片
    2. images JSON 中未被 content 引用的图片也插入末尾
    """
    import re
    import requests
    from docx.shared import Inches

    # 正则匹配 ![caption](url)
    IMG_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

    parts = IMG_RE.split(content)
    # split 返回: [text0, caption1, url1, text1, caption2, url2, text2, ...]

    for i, part in enumerate(parts):
        if i % 3 == 0:
            # 文本段
            if part.strip():
                para = doc.add_paragraph()
                run = para.add_run(part)
                run.font.name = "宋体"
                rPr = run._element.get_or_add_rPr()
                from docx.oxml.ns import qn as _qn
                rPr.rFonts.set(_qn("w:eastAsia"), "宋体")
                run.font.size = __import__('docx.shared').Pt(12)
        elif i % 3 == 2:
            # 图片 URL
            url = part
            try:
                resp = requests.get(url, timeout=10)
                if resp.status_code == 200:
                    from io import BytesIO
                    img_stream = BytesIO(resp.content)
                    _insert_image_with_caption(doc, img_stream)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f'[图片加载失败: {url}]')
                    run.font.color.rgb = __import__('docx.shared').RGBColor(0xcc, 0x00, 0x00)
            except Exception:
                para = doc.add_paragraph()
                run = para.add_run(f'[图片加载失败: {url}]')
                run.font.color.rgb = __import__('docx.shared').RGBColor(0xcc, 0x00, 0x00)

    # 插入 images JSON 中未被 content 引用的额外图片
    referenced_urls = set(IMG_RE.findall(content))
    for img in images:
        url = img.get('url', '')
        if url and url not in referenced_urls:
            try:
                resp = requests.get(url, timeout=10)
                if resp.status_code == 200:
                    from io import BytesIO
                    _insert_image_with_caption(doc, BytesIO(resp.content), caption=img.get('caption', ''))
            except Exception:
                pass


def _insert_image_with_caption(doc, img_stream, caption: str = ''):
    """插入图片并居中，可选附注"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_stream, width=Inches(3.5))
    except Exception:
        # 图片格式不支持，跳过
        return

    if caption:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True


def generate_docx(title: str, questions: list) -> bytes:
    """Generate A4 DOCX with paper formatting."""
    from docx import Document
    from docx.shared import Pt, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Watermark: header ──
    WATERMARK = "AI辅助生成｜仅供备课参考｜使用前请核对"
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(WATERMARK)
    hr.font.size = Pt(8)
    hr.font.color.rgb = __import__('docx.shared').RGBColor(0x99, 0x99, 0x99)
    hr.font.name = "宋体"
    hr.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ── Watermark: footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(WATERMARK)
    fr.font.size = Pt(8)
    fr.font.color.rgb = __import__('docx.shared').RGBColor(0x99, 0x99, 0x99)
    fr.font.name = "宋体"
    fr.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()  # spacer

    # ── Questions ──
    for q in questions:
        # Question header
        q_text = f"{q['index']}. "
        type_labels = {
            "single_choice": "【单选题】",
            "multi_choice": "【多选题】",
            "true_false": "【判断题】",
            "fill_blank": "【填空题】",
            "short_answer": "【解答题】",
        }
        q_text += type_labels.get(q["type"], "")

        para = doc.add_paragraph()
        run = para.add_run(q_text)
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = Pt(12)

        # Question content (with Markdown image support)
        content = q.get("content", "")
        images = q.get("images", [])
        _render_content_with_images(doc, content, images)

        # Options
        if q.get("options"):
            for opt in q["options"]:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1.0)
                run = para.add_run(opt)
                run.font.name = "宋体"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)

        doc.add_paragraph()  # spacer

    # ── No answer section (production: uploaded questions don't carry answers) ──

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── PDF Generation (via LibreOffice) ──────────────────────────

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    """
    Accepts the same JSON as /export/docx.
    1. Generate DOCX
    2. Convert to PDF via LibreOffice headless
    3. Return PDF bytes
    """
    try:
        data = request.get_json()
        title = data.get("title", "试卷")
        questions = data.get("questions", [])

        docx_bytes = generate_docx(title, questions)

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "input.docx")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)

            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                check=True,
                timeout=30,
            )

            pdf_path = os.path.join(tmpdir, "input.pdf")
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{title}.pdf",
        )
    except FileNotFoundError:
        return jsonify({"error": "LibreOffice not installed on this server"}), 503
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── PDF Image Extraction (PyMuPDF) ───────────────────────────

@app.route("/extract-images", methods=["POST"])
def extract_images():
    """
    从 PDF 中提取嵌入图片

    POST multipart/form-data { file: pdf }
    返回: { images: [{ pageNum, x, y, width, height, ext, cosUrl }] }
    """
    import fitz
    import uuid as _uuid

    if "file" not in request.files:
        return jsonify({"error": "请上传 PDF 文件"}), 400

    file = request.files["file"]
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "仅支持 PDF 格式"}), 400

    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        images = []

        # 确保上传目录存在
        upload_dir = os.path.join(os.getcwd(), "uploads", "images")
        os.makedirs(upload_dir, exist_ok=True)

        for page_num in range(len(doc)):
            page = doc[page_num]
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    ext = base_image["ext"]

                    # Dev: 本地存储
                    filename = f"{_uuid.uuid4()}.{ext}"
                    save_path = os.path.join(upload_dir, filename)
                    with open(save_path, "wb") as f:
                        f.write(image_bytes)
                    cos_url = f"/uploads/images/{filename}"

                    # 图片在页面中的位置
                    rects = page.get_image_rects(xref)
                    for rect in rects:
                        images.append({
                            "pageNum": page_num + 1,
                            "x": rect.x0,
                            "y": rect.y0,
                            "width": rect.width,
                            "height": rect.height,
                            "ext": ext,
                            "cosUrl": cos_url,
                        })
                except Exception:
                    pass  # 跳过无法提取的图片

        doc.close()
        return jsonify({"images": images})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Thumbnail Generation ─────────────────────────────────

@app.route("/generate-thumbnail", methods=["POST"])
def generate_thumbnail():
    """
    POST multipart/form-data { file: pdf/docx }
    返回第一页的 PNG 缩略图 (200px 宽)
    """
    import fitz
    import uuid as _uuid

    if "file" not in request.files:
        return jsonify({"error": "请上传文件"}), 400

    file = request.files["file"]
    try:
        doc_bytes = file.read()
        ext = file.filename.rsplit(".", 1)[-1].lower() if file.filename else "pdf"

        # DOCX → 通过 LibreOffice 转 PDF 再截第一页
        if ext == "docx":
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, f"input.docx")
                with open(docx_path, "wb") as f:
                    f.write(doc_bytes)
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    check=True, timeout=30,
                )
                pdf_path = os.path.join(tmpdir, "input.pdf")
                with open(pdf_path, "rb") as f:
                    doc_bytes = f.read()

        # PDF → PyMuPDF 截第一页
        pdf_doc = fitz.open(stream=doc_bytes, filetype="pdf")
        page = pdf_doc[0]
        pix = page.get_pixmap(dpi=72)
        # 缩放至 200px 宽
        scale = 200 / pix.width
        pix = page.get_pixmap(dpi=int(72 * scale))
        img_bytes = pix.tobytes("png")
        pdf_doc.close()

        # 保存
        filename = f"{_uuid.uuid4()}.png"
        upload_dir = os.path.join(os.getcwd(), "uploads", "thumbnails")
        os.makedirs(upload_dir, exist_ok=True)
        save_path = os.path.join(upload_dir, filename)
        with open(save_path, "wb") as f:
            f.write(img_bytes)

        return jsonify({"thumbnailUrl": f"/uploads/thumbnails/{filename}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Health ──

@app.route("/health")
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
